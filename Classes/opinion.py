import datetime

from  docx import Document
import os
from docx.shared import Inches, Cm, Pt
from openpyxl import Workbook
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Opinion:
    def __init__(self,info, devices):
        name = info["Nazwa pliku"]
        self.path_to_folder = './Wygenerowane_opinie'
        self.devices = devices

        months = {
            1: "styczniu",
            2: "lutym",
            3: "marcu",
            4: "kwietniu",
            5: "maju",
            6: "czerwcu",
            7: "lipcu",
            8: "sierpniu",
            9: "wrześniu",
            10: "październiku",
            11: "listopadzie",
            12: "grudniu",
        }

        self.doc = Document()
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Trebuchet MS'

        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5

        

        ### Header
        header = self.doc.sections[0].header

        table = header.add_table(rows=1, cols=3, width=Cm(16.84))
        table.autofit = False
        table.columns[0].width = Cm(2.34)
        table.columns[1].width = Cm(11.25)
        table.columns[2].width = Cm(3.25)

        for i, width in enumerate([Cm(2.34), Cm(11.25), Cm(3.25)]):
            for row in table.rows:
                row.cells[i].width = width

        table.cell(0, 0).paragraphs[0].add_run().add_picture('./Other_files/logo.png', width=Cm(1.88), height=Cm(1.35))

        run = table.cell(0, 1).paragraphs[0].clear().add_run(
            f"Opinia dot. stanu technicznego sprzętu komputerowego nr {info['Numer opinii']}/{datetime.date.today().strftime('%Y')}"
        )
        run.font.size = Pt(14)
        run.bold = True
        table.cell(0, 1).paragraphs[0].alignment = 1

        run = table.cell(0, 2).paragraphs[0].clear().add_run(f"\n\n\nData: {datetime.date.today().strftime('%d.%m.%Y')}")
        run.font.size = Pt(10)
        run.bold = False
        table.cell(0, 2).paragraphs[0].alignment = 2

        ### Document
        for_who = self.doc.add_paragraph(f"\nDla: {info['Dla kogo']}")
        for run in for_who.runs:
            run.font.size = Pt(10)
        reporter = self.doc.add_paragraph(f"Zgłaszający: {info['Zglaszajacy']}\n\n")
        for run in reporter.runs:
            run.font.size = Pt(10)

        p1 = self.doc.add_paragraph(f"\tZespół Serwisu IT przesyła informacje dot. stanu technicznego sprzętu "
                   f"dostarczonego w {months[int(datetime.date.today().strftime('%m'))]} "
                   f"{datetime.date.today().strftime('%Y')} roku do serwisu:")
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p1.runs:
            run.font.size = Pt(10)

        ### Devices
        for idx, device in enumerate(self.devices):
            # print(device.info['Data zakupu'])
            p2 = self.doc.add_paragraph(
                f"{device.info['Typ']} {device.info['Model']} "
                f"o numerach: inw. {device.info['Nr inwentarzowy']}, ID {device.info['Numer ewidencyjny']}, "
                f"SAP {device.info['Numer środka trwałego']} z "
                f"{datetime.datetime.strptime(device.info['Data zakupu'], '%d.%m.%Y').year} roku. "
                f"Po sprawdzeniu stwierdzono, że sprzęt jest sprawny technicznie i posiada standardowe ślady użytkowania. "
                f"Parametry sprzętu:",
                style="List Number"
            )
            p2.paragraph_format.left_indent = Inches(0.75)
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p2.runs:
                run.font.size = Pt(10)

            # Kolejne:
            for line in [
                f"Procesor: {device.info['Model procesora']}",
                f"Grafika: {device.info['Grafika']}",
                f"Pamięć RAM: {device.info['Pamięć RAM (GB)']}",
                f"Dysk Twardy: {device.info['Pamięć HDD (GB)']}",
                f"System operacyjny: {device.info['System operacyjny']}"
            ]:
                p = self.doc.add_paragraph(line, style="List Bullet")
                p.paragraph_format.left_indent = Inches(1)
                for run in p.runs:
                    run.font.size = Pt(10)
            formula = self.doc.add_paragraph(f"\t{device.info['Formula']}")
            formula.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in formula.runs:
                run.font.size = Pt(10)

        p3 = self.doc.add_paragraph(f"\nOpinię sporządził: {info['Autor']}")
        for run in p3.runs:
            run.font.size = Pt(10)
        p4 = self.doc.add_paragraph(f"\t\t\t\t\t\t\t\tZ poważaniem,")
        for run in p4.runs:
            run.font.size = Pt(10)

        def add_page_number(paragraph, font_size_pt):
            run = paragraph.add_run("Strona ")
            run.font.size = Pt(font_size_pt)

            # Pole PAGE (aktualna strona)
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText1 = OxmlElement('w:instrText')
            instrText1.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            r_element = run._r
            r_element.append(fldChar1)
            r_element.append(instrText1)
            r_element.append(fldChar2)
            r_element.append(fldChar3)

            # Dodaj tekst " z "
            run2 = paragraph.add_run(" z ")
            run2.font.size = Pt(font_size_pt)

            # Pole NUMPAGES (liczba wszystkich stron)
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'begin')

            instrText2 = OxmlElement('w:instrText')
            instrText2.text = "NUMPAGES"

            fldChar5 = OxmlElement('w:fldChar')
            fldChar5.set(qn('w:fldCharType'), 'separate')

            fldChar6 = OxmlElement('w:fldChar')
            fldChar6.set(qn('w:fldCharType'), 'end')

            run3 = paragraph.add_run()
            run3.font.size = Pt(font_size_pt)
            r_element2 = run3._r
            r_element2.append(fldChar4)
            r_element2.append(instrText2)
            r_element2.append(fldChar5)
            r_element2.append(fldChar6)

        def set_cell_border(cell, top="single", size="6", color="000000"):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)

            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), top)
            top_border.set(qn('w:sz'), size)
            top_border.set(qn('w:space'), "0")
            top_border.set(qn('w:color'), color)
            tcBorders.append(top_border)

        ### Footer
        footer = self.doc.sections[0].footer
        ftable = footer.add_table(rows=1, cols=3, width=Cm(17.44))
        ftable.autofit = False
        ftable.allow_autofit = False
        for row in ftable.rows:
            row.cells[0].width = Cm(8.75)
            row.cells[1].width = Cm(3.84)
            row.cells[2].width = Cm(4.85)

        para = ftable.cell(0, 0).paragraphs[0]
        para.clear()

        run1 = para.add_run("Grupa Azoty S.A.")
        run1.font.size = Pt(9)
        run1.bold = False

        run2 = para.add_run("\nDepartament Korporacyjny Informatyki/Zespół Serwisu IT")
        run2.font.size = Pt(8)
        run2.bold = False

        right_cell = ftable.cell(0, 2).paragraphs[0]
        right_cell.alignment = 2
        add_page_number(right_cell, 9)

        # Obramowanie górne tylko w wybranych komórkach
        set_cell_border(ftable.cell(0, 0))  # Top border
        set_cell_border(ftable.cell(0, 1))  # Top border
        set_cell_border(ftable.cell(0, 2))  # Top border

        os.mkdir(self.path_to_folder + "/" + name)
        self.path_to_opinion = self.path_to_folder + "/" + name +'/'
        self.doc.save(f"{self.path_to_opinion}Opinia {info['Numer opinii']}.docx")